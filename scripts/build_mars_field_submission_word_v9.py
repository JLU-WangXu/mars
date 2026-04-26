from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "paper_bundle_v1" / "MARS_FIELD_Submission_Style_v9.docx"
FIG = ROOT / "outputs" / "paper_bundle_v1" / "figures"

CLAIMS = ROOT / "docs" / "mars_field_core_claims_v2.md"
INTRO_DISC = ROOT / "docs" / "mars_field_intro_discussion_draft_v2.md"
RELATED = ROOT / "docs" / "mars_field_related_work_full_v3.md"
METHODS = ROOT / "docs" / "mars_field_methods_full_v3.md"
LEGENDS = ROOT / "docs" / "mars_field_figure_legends_v3.md"


def set_font(run, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_p(doc_or_cell, text: str, *, size: float = 10.5, bold: bool = False, italic: bool = False, align=None, color: str | None = None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.space_after = Pt(4)
    return p


def add_h(doc_or_cell, text: str, level: int = 1):
    size = {1: 14.0, 2: 11.8, 3: 10.6}[level]
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(3)


def add_figure(doc, image_name: str, caption: str, width: float = 6.0) -> None:
    path = FIG / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = c.add_run(caption)
    set_font(run, size=8.3, italic=True)
    c.paragraph_format.space_after = Pt(6)


def emit_markdown(doc: Document, path: Path) -> None:
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            add_h(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            add_h(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_h(doc, line[4:].strip(), 3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:].strip())
            set_font(run, size=10.0)
            p.paragraph_format.line_spacing = 1.06
            p.paragraph_format.space_after = Pt(2)
        elif line[:2].isdigit() and line[2:4] == ". ":
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(line[4:].strip())
            set_font(run, size=10.0)
            p.paragraph_format.line_spacing = 1.06
            p.paragraph_format.space_after = Pt(2)
        else:
            add_p(doc, line, size=10.3)


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    add_p(doc, "MARS-FIELD", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color="1F4E79")
    add_p(doc, "A unified evidence-to-sequence controller for protein engineering", size=12.8, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Submission-style draft v9", size=10.2, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_h(doc, "Article Spine", 1)
    emit_markdown(doc, CLAIMS)

    add_h(doc, "Main Figures", 1)
    add_figure(doc, "figure1_mars_field_architecture_v2.png", "Figure 1 | MARS-FIELD integrates heterogeneous evidence into a shared residue field.", width=6.05)
    add_figure(doc, "figure2_benchmark_claim_v5.png", "Figure 2 | Benchmark claim figure.", width=6.05)
    add_figure(doc, "figure3_mechanism_limitations_v5.png", "Figure 3 | Mechanism and limitation figure.", width=6.05)
    add_figure(doc, "figure4_case_studies_master_v3.png", "Figure 4 | PSE-derived case-study composite.", width=6.05)

    add_h(doc, "Introduction and Discussion", 1)
    emit_markdown(doc, INTRO_DISC)

    add_h(doc, "Related Work", 1)
    emit_markdown(doc, RELATED)

    add_h(doc, "Methods", 1)
    emit_markdown(doc, METHODS)

    add_h(doc, "Figure Legends", 1)
    emit_markdown(doc, LEGENDS)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved submission-style Word draft to {OUT}")


if __name__ == "__main__":
    main()
