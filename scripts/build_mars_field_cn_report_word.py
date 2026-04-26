from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "mars_field_technical_report_cn_v1.md"
OUT = ROOT / "outputs" / "paper_bundle_v1" / "MARS_FIELD_中文技术报告_v1.docx"


def set_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


def add_text(doc: Document, text: str, *, size: float = 10.5, bold: bool = False, align=None, color: str | None = None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    return p


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    add_text(doc, "MARS-FIELD 中文技术报告", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color="1F4E79")
    add_text(doc, "内部技术说明与论文准备辅助文档", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    for raw in SRC.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            add_text(doc, line[2:].strip(), size=15, bold=True)
        elif line.startswith("## "):
            add_text(doc, line[3:].strip(), size=12.5, bold=True)
        elif line.startswith("### "):
            add_text(doc, line[4:].strip(), size=11.2, bold=True)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:].strip())
            set_font(run, size=10.2)
            p.paragraph_format.line_spacing = 1.08
            p.paragraph_format.space_after = Pt(2)
        else:
            add_text(doc, line, size=10.5)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved Chinese technical report to {OUT}")


if __name__ == "__main__":
    main()
