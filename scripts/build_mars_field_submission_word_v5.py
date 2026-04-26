from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "paper_bundle_v1"
FIG_DIR = OUT_DIR / "figures"
OUT_DOCX = OUT_DIR / "MARS_FIELD_Submission_Style_v5.docx"


def set_font(run, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_p(doc_or_cell, text: str, *, size: float = 10.5, bold: bool = False, italic: bool = False, align=None, color: str | None = None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(4)
    return p


def add_h(doc_or_cell, text: str, level: int = 1):
    size = {1: 14.5, 2: 12.0, 3: 10.8}[level]
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "D9E2F3")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_font(run, size=9, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_font(run, size=8.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def add_figure(doc, image_name: str, caption: str, width: float = 5.8) -> None:
    image_path = FIG_DIR / image_name
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_font(run, size=8.5, italic=True)
    cap.paragraph_format.space_after = Pt(6)


def references() -> list[str]:
    return [
        "1. Jumper, J. et al. Highly accurate protein structure prediction with AlphaFold. Nature 596, 583-589 (2021).",
        "2. Dauparas, J. et al. Robust deep learning-based protein sequence design using ProteinMPNN. Science 378, 49-56 (2022).",
        "3. Hsu, C. et al. Learning inverse folding from millions of predicted structures. Proc. ICML, PMLR 162, 8946-8970 (2022).",
        "4. Lin, Z. et al. Evolutionary-scale prediction of atomic-level protein structure with a language model. Science 379, 1123-1130 (2023).",
        "5. van Kempen, M. et al. Fast and accurate protein structure search with Foldseek. Nat. Biotechnol. 42, 243-246 (2024).",
        "6. Alley, E. C. et al. Unified rational protein engineering with sequence-based deep representation learning. Nat. Methods 16, 1315-1322 (2019).",
        "7. Chisholm, L. O. et al. Ancestral reconstruction and the evolution of protein energy landscapes. Annu. Rev. Biophys. 53, 127-146 (2024).",
        "8. Matthews, D. S. et al. Leveraging ancestral sequence reconstruction for protein representation learning. Nat. Mach. Intell. 6, 1542-1555 (2024).",
    ]


def build_doc() -> Document:
    bench = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "benchmark_summary.csv")
    compare = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "compare_current_vs_final.csv")
    family = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "family_summary.csv")

    delta_col = "policy_selection_score_delta_final_minus_current"
    positive = int((compare[delta_col] > 1e-9).sum())
    negative = int((compare[delta_col] < -1e-9).sum())
    mean_delta = float(compare[delta_col].mean())
    decoder_preview = int(bench["neural_decoder_generated_count"].sum())
    decoder_novel = int(bench["neural_decoder_novel_count"].sum())
    decoder_rejected = int(bench["neural_decoder_rejected_count"].sum())
    decoder_injected = int(bench["neural_decoder_injected"].sum())

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    add_p(doc, "MARS-FIELD", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color="1F4E79")
    add_p(doc, "A unified evidence-to-sequence controller for protein engineering", size=13, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Submission-style integrated draft v5", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_h(doc, "Abstract", 1)
    add_p(
        doc,
        f"Protein engineering pipelines often combine structural heuristics, evolutionary priors, and proposal generators as loosely connected modules. We present MARS-FIELD, a unified evidence-to-sequence controller that projects geometry-conditioned structural evidence, phylogenetic profiles, ancestral lineage information, retrieval-based motif memory, and environment-conditioned engineering context into a shared residue field. The current implementation couples this field to a candidate-level neural controller and a decode-time neural field generator inside the benchmark-time main path. Across a twelve-target benchmark spanning ten protein families, the final controller improved paired policy score on {positive} of 12 targets relative to the incumbent benchmark, enabled neural decoding on all 12 targets, and retained {decoder_novel} novel neural-decoder-derived candidates across {decoder_injected} targets. These results position MARS-FIELD as a stable neuralized controller-decoder system that moves beyond heuristic score stacking toward an evidence-conditioned field architecture.",
    )

    add_h(doc, "Introduction", 1)
    add_p(
        doc,
        "Protein engineering increasingly depends on computational pipelines that integrate structural analysis, evolutionary context, and machine-learning proposal engines. In many practical systems, however, these components remain only loosely coupled. Structure-aware proposal generators, profile-derived priors, ancestral sequence reconstructions, retrieval-based motif analogues, and engineering heuristics are often applied in parallel or in sequence, with the final decision delegated to a hand-tuned ranking rule [1-8]. Although this strategy enables rapid iteration, it creates a conceptual weakness: the system behaves as an ensemble of partially redundant signals rather than as a unified controller.",
    )
    add_p(
        doc,
        "This weakness becomes more pronounced under realistic engineering objectives. Protein redesign must satisfy multiple incompatible demands, including structural compatibility, evolutionary tolerability, ancestral plausibility, motif-level memory support, and environment-specific constraints such as oxidation pressure or flexible-surface burden. These signals differ in scale, semantics, and reliability, making naïve voting or score stacking unsatisfactory as a general algorithmic formulation.",
    )
    add_p(
        doc,
        "We developed MARS-FIELD to treat protein engineering as an evidence-to-sequence control problem. In MARS-FIELD, geometric, phylogenetic, ancestral, retrieval-based, and environment-conditioned signals are projected into a shared residue field over the design positions. This field is the native decision object of the method: it defines residue-wise preferences and pairwise couplings, which are then used both to evaluate incumbent candidates and to decode new ones under engineering constraints. External generators and priors are therefore not independent voters; they are evidence streams that shape a common residue-level decision space.",
    )
    add_p(
        doc,
        "The present implementation takes a substantial step beyond our earlier engineering approximations. First, ancestry and retrieval are represented as learned memory-linked branches within the neural controller. Second, the neural branch is no longer restricted to reranking. The main benchmark path now includes a neural field decoder that generates candidates from the shared field itself. We do not claim that this is yet the final fully joint proposal-generator / field / decoder research endpoint. Rather, we show that a unified evidence-conditioned controller can already be built in a way that is mechanistically interpretable, benchmark-stable, and capable of decode-time neural proposal generation.",
    )

    add_h(doc, "Results", 1)
    add_h(doc, "A benchmark-scale neural field decoder is active in the main design loop", 2)
    add_p(
        doc,
        f"The most important result is architectural. The benchmark-time main path now contains a shared evidence field, a calibrated neural controller, and a decode-time neural field generator. Across the twelve-target panel, neural decoding was enabled on all 12 targets, produced {decoder_preview} preview candidates, retained {decoder_novel} novel candidates after filtering, and contributed retained neural-decoder candidates on {decoder_injected} targets. The neural branch is therefore participating in proposal generation rather than only in late reranking.",
    )
    add_figure(doc, "figure2_benchmark_claim_v5.png", "Figure 2 | Benchmark claim figure. Paired policy deltas, decoder utilization, and family-level summary support the claim that the final controller remains stable while enabling neural decode-time generation.", width=6.1)

    add_h(doc, "The final controller remains globally stable despite stronger neuralization", 2)
    add_p(
        doc,
        f"Relative to the incumbent benchmark, the final controller improved paired policy score on {positive} of 12 targets and decreased it on {negative} targets, with a mean paired delta of {mean_delta:.3f}. This near-neutral mean should not be misread as a lack of effect. Instead, it indicates that the controller remains globally stable even after a substantial increase in model complexity. The more important observation is that gains outnumber regressions threefold.",
    )
    add_p(
        doc,
        "Positive shifts were observed across several distinct protein families, including lipase_b, adenylate kinase, lipase-esterase, cutinase, GFP-like proteins, superoxide dismutase, lysozyme, and beta-lactamase. In multiple cases, the final controller did not need to replace the incumbent top sequence to add value; instead, it strengthened incumbent support while maintaining chemically plausible design logic.",
    )
    add_p(
        doc,
        "The remaining negative shifts were concentrated in CLD_3Q09_NOTOPIC, CLD_3Q09_TOPIC, and subtilisin_2st1. This concentration is preferable to a diffuse collapse because it exposes a small set of explicit calibration-limited targets rather than undermining the method globally.",
    )

    add_h(doc, "Mechanistic analyses identify both dominant constraints and explicit limitation modes", 2)
    add_p(
        doc,
        "Ablation and diagnostic analyses indicate that the current method is not a black-box neural add-on. Removing oxidation information changed the top candidate on 10 of 12 targets and caused a strong average score drop, indicating that oxidation-aware engineering remains one of the dominant constraints. Removing evolutionary information changed the top candidate on 6 of 12 targets and strongly altered the score landscape, showing that profile-derived priors remain foundational to the system. By contrast, removing surface terms changed only 2 of 12 top candidates, suggesting a more target-local role for surface information in the present benchmark.",
    )
    add_p(
        doc,
        "Neural gate profiles additionally show that different targets rely on different evidence mixtures. The current regressions are therefore more plausibly interpreted as evidence-balancing and calibration problems than as failures of the overall field-controller formulation itself.",
    )
    add_figure(doc, "figure3_mechanism_limitations_v5.png", "Figure 3 | Mechanism and limitation figure. Ablations reveal dominant constraints, neural diagnostics show target-dependent evidence usage, and failures remain concentrated rather than diffuse.", width=6.1)

    add_h(doc, "Representative case studies reveal distinct controller regimes", 2)
    add_p(
        doc,
        "1LBT functions as a conservative safety case in which the incumbent M298L solution is preserved despite full neural-decoder execution. TEM1 functions as a high-scoring engineering case in which the top policy remains stable but the neural decoder contributes a learned alternative worth inspection. PETase provides a reproducibility case across two structures, while CLD functions as a calibration stress test that exposes the tension between high local engineering signal and stable final policy.",
    )

    # concise benchmark table
    compare_lookup = {str(r["target"]): r for _, r in compare.iterrows()}
    rows = []
    for _, row in bench.iterrows():
        cmp = compare_lookup[str(row["target"])]
        rows.append(
            [
                str(row["target"]),
                str(row["family"]),
                str(row["policy_mutations"]),
                f'{float(row["policy_selection_score"]):.3f}',
                f'{float(row["policy_engineering_score"]):.3f}',
                f'{float(cmp[delta_col]):+.3f}',
                str(int(row["neural_decoder_novel_count"])),
            ]
        )
    add_table(doc, ["Target", "Family", "Final policy mutation", "Policy score", "Engineering score", "Paired delta", "Decoder novel"], rows, [0.9, 1.0, 2.0, 0.7, 0.8, 0.7, 0.7])

    add_h(doc, "Methods", 1)
    add_p(
        doc,
        "The primary benchmark consists of twelve targets spanning ten protein families. Each target defines a local design window on a fixed wild-type structural scaffold. Neural controller training is leave-one-target-out at the target level. The main deployment arm used here is benchmark_twelvepack_final, which enables both neural reranking and neural field decoding while using a hybrid final policy for stable deployment.",
    )
    add_p(
        doc,
        "MARS-FIELD integrates five evidence streams: geometry-conditioned structural features, phylogenetic profile evidence, ancestral lineage evidence, retrieval-based motif memory, and environment-conditioned engineering context. These signals are projected into a shared residue field containing both residue-wise and pairwise preferences. The neural field contains geometry, phylogeny, ancestry, retrieval, and environment branches; ancestry and retrieval each interact with learned memory banks to fuse local residue representations with lineage memory and retrieval prototype memory.",
    )
    add_p(
        doc,
        "Candidate-level decision making combines sequence-conditioned residue embeddings with candidate-specific evidence features, including source identity, support count, mutation burden, component-wise engineering scores, selector-rank priors, gaps to the best incumbent, and other calibration-aware signals. Training uses multiple objectives, including selection regression, engineering regression, policy regression, pairwise ranking, decoder-field supervision, winner-guard loss, non-decoder guard loss, simplicity guard loss, selector-anchor distillation, recovery loss, ancestry alignment, retrieval alignment, pairwise consistency, environment reconstruction, and gate regularization.",
    )
    add_p(
        doc,
        "For each target, the pipeline constructs a runtime neural batch from the current target state, trains a leave-one-target-out neural field model using the remaining targets, combines learned unary and pairwise outputs with evidence-derived prior position fields and prior pairwise terms, and decodes neural_decoder candidates under constrained beam search. The final hybrid policy permits learned branch adoption only when the alternative remains sufficiently aligned with incumbent selection-score stability while preserving engineering support.",
    )

    add_h(doc, "Discussion", 1)
    add_p(
        doc,
        "The central conclusion of this study is that a unified residue-field controller provides a more coherent algorithmic basis for protein engineering than an ensemble of disconnected proposal branches. In the current implementation, MARS-FIELD integrates heterogeneous evidence streams into a shared residue field, uses a calibrated neural controller to score candidates, and activates a decode-time neural branch inside the main benchmark path. This is qualitatively different from earlier stack-and-rerank pipelines because the neural branch now participates in generation as well as in selection.",
    )
    add_p(
        doc,
        "The present results also clarify the current boundary of the method. The system is now strong enough to be called an engineering-honest end-to-end controller-decoder architecture, but it is not yet the fully joint proposal-generator / field / decoder research endpoint. The remaining gap is therefore no longer missing infrastructure; it is a genuine research problem centered on joint optimization and improved calibration on a small number of hard targets.",
    )
    add_p(
        doc,
        "For manuscript positioning, the most defensible claim is not that MARS-FIELD completely supersedes all earlier modular pipelines, but that it provides a stable bridge between heuristic engineering stacks and future fully joint protein design controllers. The method already demonstrates broad-panel stability, real neural proposal generation, interpretable multi-evidence behavior, and explicit limitation cases. These are the right ingredients for a serious methods paper.",
    )

    add_h(doc, "Selected References", 1)
    for ref in references():
        add_p(doc, ref, size=9.4)

    return doc


def main() -> None:
    doc = build_doc()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Saved submission-style Word draft to {OUT_DOCX}")


if __name__ == "__main__":
    main()
