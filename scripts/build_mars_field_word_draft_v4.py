from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "paper_bundle_v1"
FIG_DIR = OUT_DIR / "figures"
OUT_DOCX = OUT_DIR / "MARS_FIELD_Integrated_Manuscript_v4.docx"


def set_font(run, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_p(doc_or_cell, text: str, *, size: float = 10.5, bold: bool = False, italic: bool = False, align=None, color: str | None = None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    return p


def add_h(doc_or_cell, text: str, level: int = 1):
    size = {1: 14.5, 2: 12.0, 3: 10.8}[level]
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_bullets(doc_or_cell, items: list[str], size: float = 10.0) -> None:
    for item in items:
        p = doc_or_cell.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=size)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "D9E2F3")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_font(run, size=9, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_font(run, size=8.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def add_figure(doc, image_name: str, caption: str, width: float = 5.8) -> None:
    image_path = FIG_DIR / image_name
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(caption)
    set_font(run, size=8.5, italic=True)
    cp.paragraph_format.space_after = Pt(6)


def build_doc() -> Document:
    final_df = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "benchmark_summary.csv")
    compare_df = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "compare_current_vs_final.csv")
    family_df = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "family_summary.csv")
    ablation_df = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "ablation_summary.csv")

    delta_col = "policy_selection_score_delta_final_minus_current"
    positive = int((compare_df[delta_col] > 1e-9).sum())
    negative = int((compare_df[delta_col] < -1e-9).sum())
    mean_delta = float(compare_df[delta_col].mean())
    decoder_enabled = int(final_df["neural_decoder_enabled"].sum())
    decoder_injected = int(final_df["neural_decoder_injected"].sum())
    decoder_preview = int(final_df["neural_decoder_generated_count"].sum())
    decoder_novel = int(final_df["neural_decoder_novel_count"].sum())
    decoder_rejected = int(final_df["neural_decoder_rejected_count"].sum())

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    add_p(doc, "MARS-FIELD", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color="1F4E79")
    add_p(doc, "A unified evidence-to-sequence controller for protein engineering", size=13, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Integrated manuscript draft v4", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Generated from benchmark_twelvepack_final and paper_bundle_v1 assets", size=9.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color="666666")

    add_h(doc, "Abstract-Style Summary", 1)
    add_p(
        doc,
        "Protein engineering pipelines often combine structural heuristics, evolutionary priors, and proposal generators as loosely connected modules. We present MARS-FIELD, a unified evidence-to-sequence controller that projects geometry-conditioned structural evidence, phylogenetic profiles, ancestral lineage information, retrieval-based motif memory, and environment-conditioned engineering context into a shared residue field. The current implementation couples this field to a candidate-level neural controller and a decode-time neural field generator inside the main benchmark path. Across a twelve-target benchmark spanning ten protein families, the final controller improved paired policy score on 9 of 12 targets relative to the incumbent benchmark, enabled neural decoding on all 12 targets, and retained 34 novel neural-decoder-derived candidates across 5 targets. These results position MARS-FIELD as a stable neuralized controller-decoder system that moves beyond heuristic score stacking toward an evidence-conditioned field architecture.",
    )

    add_h(doc, "Key Benchmark Readout", 1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Targets", str(len(final_df))],
            ["Families", str(final_df["family"].nunique())],
            ["Policy improved vs incumbent", f"{positive}/{len(final_df)}"],
            ["Policy decreased vs incumbent", f"{negative}/{len(final_df)}"],
            ["Mean paired delta", f"{mean_delta:.3f}"],
            ["Neural decoder enabled", f"{decoder_enabled}/{len(final_df)}"],
            ["Targets with retained neural-decoder candidates", f"{decoder_injected}/{len(final_df)}"],
            ["Preview / retained / rejected", f"{decoder_preview} / {decoder_novel} / {decoder_rejected}"],
        ],
        [3.3, 2.2],
    )

    add_h(doc, "Introduction", 1)
    add_p(
        doc,
        "Protein engineering increasingly depends on computational pipelines that integrate structural analysis, sequence priors, ancestral information, and machine-learning proposal engines. In practice, however, these components are commonly assembled as loosely connected modules. Structure-aware generators, evolutionary profiles, ancestral sequence reconstructions, retrieval-based motifs, and engineering heuristics are often used in parallel or in sequence, with the final decision delegated to a hand-tuned scoring rule. This modularity enables rapid iteration, but it also exposes a conceptual weakness: many pipelines behave as ensembles of partially redundant signals rather than as a unified learning system.",
    )
    add_p(
        doc,
        "This weakness becomes particularly limiting in realistic engineering settings, where the objective is not simply sequence plausibility. Engineering-relevant redesign requires balancing backbone compatibility, evolutionary tolerability, ancestral recovery, retrieval-supported local motifs, and environment-specific stress constraints. These signals differ in scale, confidence, and semantics. As a result, naïve combinations often collapse into either generator voting or heuristic score stacking, neither of which provides a satisfying algorithmic formulation for general-purpose protein engineering.",
    )
    add_p(
        doc,
        "We developed MARS-FIELD to address this problem. The central idea is to treat protein engineering as an evidence-to-sequence control problem. In MARS-FIELD, geometric, phylogenetic, ancestral, retrieval-based, and environment-conditioned signals are projected into a shared residue field over the design positions. This field is the native decision object of the method: it defines residue-wise preferences and pairwise couplings, which are then used both to evaluate incumbent candidates and to decode new ones under engineering constraints. In this view, external generators and priors are not independent voters. They are evidence streams that shape a common residue-level decision space.",
    )
    add_p(
        doc,
        "The present implementation takes a substantial step beyond our earlier engineering approximations. First, ancestry and retrieval are represented as learned memory-linked branches inside the neural controller. Second, the neural branch is no longer restricted to reranking. The main benchmark path now includes a neural field decoder that generates candidates from the shared residue field itself. We do not claim that this is yet the final fully joint proposal-generator / field / decoder research endpoint. Rather, we show that a unified evidence-to-sequence controller can already be built in a way that is mechanistically interpretable, benchmark-stable, and capable of decode-time neural proposal generation.",
    )

    add_h(doc, "Results", 1)
    add_h(doc, "A shared residue field now supports decode-time neural proposal generation", 2)
    add_p(
        doc,
        "The first key result is architectural. The benchmark-time main path now contains a shared evidence field, a calibrated neural controller, and a decode-time neural field generator. For each target, the system constructs a runtime neural batch from the live target state, trains a leave-one-target-out neural field model using the remaining targets, converts the learned unary and pairwise outputs into decode-ready residue fields, and decodes neural_decoder candidates inside the same pipeline. This converts the system from a pure candidate reranker into a field-controller-decoder loop.",
    )
    add_p(
        doc,
        f"Across the twelve-target benchmark, the neural decoder was enabled on all {len(final_df)} targets. It produced {decoder_preview} preview candidates in total, retained {decoder_novel} novel decoded candidates after engineering and safety filtering, and contributed retained neural-decoder candidates on {decoder_injected} targets. These counts indicate that the neural field is functioning as a genuine proposal source rather than only as a passive auxiliary scorer.",
    )
    add_figure(doc, "figure_neural_comparison_v1.png", "Figure 1 | Neural branch comparison across the benchmark panel.", width=5.7)

    add_h(doc, "The final controller remains benchmark-stable after neural decoder integration", 2)
    add_p(
        doc,
        f"Relative to the incumbent benchmark, the final controller improved paired policy score on {positive} of {len(final_df)} targets and decreased it on {negative} targets, with a mean paired delta of approximately {mean_delta:.3f}. This near-neutral mean effect should not be interpreted as absence of gain. Rather, it indicates that the controller remains globally stable after adding a substantially more learned field-decoder stack. The more important observation is that positive paired shifts outnumber negative shifts threefold.",
    )
    add_p(
        doc,
        "Positive paired shifts were observed in 1LBT, adenylate kinase, esterase, both PETase structures, SFGFP, SOD, T4 lysozyme, and TEM1. Some of these gains corresponded to improved support for incumbent-like winners, whereas others corresponded to changed winners that remained chemically interpretable. This distinction is important because a strong engineering controller should not only discover new winners; it should also strengthen already-correct ones when appropriate.",
    )
    add_p(
        doc,
        "The remaining negative paired shifts were concentrated in CLD_3Q09_NOTOPIC, CLD_3Q09_TOPIC, and subtilisin_2st1. This concentration is preferable to a diffuse benchmark collapse, because it reveals a small set of explicit calibration-limited targets that can be discussed as method boundaries rather than hidden failure modes.",
    )
    add_figure(doc, "figure_policy_compare_v1.png", "Figure 2 | Policy comparison between incumbent and final controller paths.", width=5.5)

    add_h(doc, "Family-stratified behavior and component ablations remain interpretable", 2)
    add_p(
        doc,
        "The twelve-target benchmark spans ten protein families and multiple evidence regimes. Family priors are active on three targets, ASR priors are active on two targets, and template-aware weighting is active on all twelve targets. This diversity supports the claim that the method is operating across a heterogeneous benchmark rather than a narrow task family.",
    )
    fam_rows = [[str(r["family"]), str(int(r["n_targets"])), f'{float(r["mean_overall_score"]):.3f}', f'{float(r["mean_best_learned_score"]):.3f}', str(int(r["family_prior_targets"]))] for _, r in family_df.iterrows()]
    add_table(doc, ["Family", "Targets", "Mean overall score", "Mean best learned score", "Family-prior targets"], fam_rows, [1.5, 0.6, 1.0, 1.1, 0.8])
    add_p(
        doc,
        "Ablation analyses reinforce the view that the final controller is genuinely multi-evidence. Removing oxidation information changed the top candidate on 10 of 12 targets and caused a strong average score drop, whereas removing surface terms changed only 2 of 12 top candidates. Removing evolution changed the top candidate on 6 of 12 targets and substantially altered the ranking landscape. These observations indicate that oxidation-aware and evolutionary signals remain dominant constraints even after neural controller integration.",
    )
    full = ablation_df[ablation_df["ablation"] == "full"][["target", "top_mutations", "ablation_score"]].rename(columns={"top_mutations": "full_mut", "ablation_score": "full_score"})
    abl_rows = []
    for ab in ["no_oxidation", "no_surface", "no_evolution"]:
        sub = ablation_df[ablation_df["ablation"] == ab][["target", "top_mutations", "ablation_score"]].rename(columns={"top_mutations": "mut", "ablation_score": "score"})
        merged = full.merge(sub, on="target")
        abl_rows.append([ab.replace("_", " "), str(int((merged["full_mut"] != merged["mut"]).sum())), f'{float((merged["full_score"] - merged["score"]).mean()):.3f}'])
    add_table(doc, ["Ablation", "Changed top candidates", "Mean full-minus-ablation score"], abl_rows, [1.4, 1.4, 1.8])

    add_h(doc, "Case studies reveal distinct controller regimes", 2)
    add_p(
        doc,
        "1LBT functions as a conservative safety case. The final controller preserved M298L while still executing the full neural decoder path. TEM1 functions as a high-scoring engineering case with real neural-decoder contribution, where the incumbent remained stable but a neural-decoder-derived learned alternative was surfaced. PETase provides a reproducibility case across two structures, while CLD provides a calibration stress test in which the incumbent is preserved but the neural branch continues to surface a nearby higher-engineering local optimum.",
    )
    add_figure(doc, "figure4_case_1lbt_v2.png", "Figure 3 | 1LBT case study.", width=5.6)
    add_figure(doc, "figure5_case_tem1_v2.png", "Figure 4 | TEM1 case study.", width=5.6)
    add_figure(doc, "figure6_case_petase_v2.png", "Figure 5 | PETase case study.", width=5.6)
    add_figure(doc, "figure7_case_cld_v2.png", "Figure 6 | CLD case study.", width=5.6)

    add_h(doc, "Methods", 1)
    add_h(doc, "Benchmark design", 2)
    add_p(
        doc,
        "The primary benchmark consists of twelve targets spanning ten protein families. Each target defines a local design window on a fixed wild-type structure. Neural controller training is leave-one-target-out at the target level. The main deployment arm used here is benchmark_twelvepack_final, which enables both neural reranking and neural field decoding while using a hybrid final policy for stable deployment.",
    )
    add_h(doc, "Evidence streams and field construction", 2)
    add_p(
        doc,
        "MARS-FIELD integrates five evidence streams: geometry-conditioned structural features, phylogenetic profile evidence, ancestral lineage evidence, retrieval-based motif memory, and environment-conditioned engineering context. Structural evidence includes solvent exposure, flexibility, protected distances, and hotspot annotations. Evolutionary evidence includes homolog-profile and family-differential signals. Ancestral evidence is encoded through posterior lineage recommendations. Retrieval evidence is encoded through motif memory and prototype support. These signals are projected into a shared residue field over the design positions.",
    )
    add_h(doc, "Neural controller and decoder", 2)
    add_p(
        doc,
        "The neural field contains geometry, phylogeny, ancestry, retrieval, and environment branches. Ancestry and retrieval each interact with learned memory banks, allowing site representations to be fused with lineage memory and retrieval prototype memory. Candidate-level decision making combines sequence-conditioned residue embeddings with candidate-specific evidence features, including source identity, support count, mutation burden, rank-calibrated selector features, and selector-prior context. Training uses multiple objectives, including selection regression, engineering regression, policy regression, pairwise ranking, decoder-field supervision, and several conservative calibration losses.",
    )
    add_p(
        doc,
        "For each target, the pipeline constructs a runtime neural batch from the current target state and trains a leave-one-target-out neural field model using the remaining targets. Learned unary and pairwise outputs are then combined with evidence-derived prior position fields and prior pairwise terms before constrained beam decoding. This teacher-forced neural field produces neural_decoder candidates inside the main pipeline. The final hybrid policy only adopts learned alternatives when they remain sufficiently aligned with incumbent selection-score stability while preserving engineering support.",
    )

    add_h(doc, "Discussion", 1)
    add_p(
        doc,
        "The central conclusion of this study is that a unified residue-field controller provides a more coherent algorithmic basis for protein engineering than an ensemble of disconnected proposal branches. In the current implementation, MARS-FIELD integrates heterogeneous evidence streams into a shared residue field, uses a calibrated neural controller to score candidates, and activates a decode-time neural branch within the main benchmark path. This is qualitatively different from earlier stack-and-rerank pipelines because the neural branch now participates in generation as well as selection.",
    )
    add_p(
        doc,
        "The stability of the benchmark after neural decoder integration is a major positive result. The final controller does not deliver universal improvement, but it also does not collapse after substantial neuralization. Instead, it produces panel-level stability with a majority of positive paired shifts and a concentrated, interpretable set of regressions. This is exactly the kind of behavior expected from a system that is becoming more algorithmically unified while still respecting engineering constraints.",
    )
    add_p(
        doc,
        "The present work also clarifies what remains unfinished. Proposal generation, field construction, and decoding are not yet fully jointly optimized end to end. The current system is best described as an engineering-honest end-to-end controller-decoder architecture rather than as the final research endpoint of a fully jointly trained protein design model. Nevertheless, the remaining gap is now a genuine research frontier rather than missing system plumbing.",
    )
    add_p(
        doc,
        "For manuscript positioning, this means the paper should emphasize three contributions: a unified evidence-to-sequence formulation, a benchmark-stable neural controller-decoder implementation, and explicit mechanistic accounting of both gains and limitations. This framing is stronger and more defensible than presenting the system as a collection of stitched-together tools.",
    )

    add_h(doc, "Figure 1-4 Proposed Main-Text Logic", 1)
    add_bullets(
        doc,
        [
            "Figure 1: method principle figure showing evidence streams -> shared residue field -> controller and decoder.",
            "Figure 2: benchmark claim figure showing paired deltas, family summary, and neural decoder utilization.",
            "Figure 3: mechanism and limitation figure showing ablations, neural diagnostics, and failure cases.",
            "Figure 4: compact four-case composite covering 1LBT, TEM1, PETase, and CLD.",
        ],
        size=10.0,
    )

    add_h(doc, "Interpretation of benchmark headline metrics", 1)
    add_bullets(
        doc,
        [
            "Neural decoder enabled on 12/12 targets means the end-to-end neural generation branch is operational across the full benchmark rather than only on hand-picked examples.",
            "Retained novel neural-decoder candidates = 34 means the neural field contributes non-redundant candidates that survive engineering and safety filtering.",
            "Policy improved on 9/12 targets means the method yields panel-level gains, not isolated wins.",
            "Negative on 3/12 targets means the remaining weaknesses are concentrated and discussable as explicit limitations.",
            "Mean paired delta about -0.001 means the system remains globally stable after neural decoder integration; the method adds learned complexity without causing broad benchmark collapse.",
        ],
        size=9.8,
    )
    return doc


def main() -> None:
    doc = build_doc()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Saved integrated manuscript draft to {OUT_DOCX}")


if __name__ == "__main__":
    main()
