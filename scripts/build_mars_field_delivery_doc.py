from __future__ import annotations

import subprocess
from datetime import datetime
import html
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_mars_field_html_delivery as delivery


ROOT = Path(__file__).resolve().parents[2]
MARS_ROOT = ROOT / "mars_stack"
OUT_DIR = ROOT / "release_packages"
DOCX_PATH = OUT_DIR / "mars_field_delivery_report_final.docx"
PDF_PATH = OUT_DIR / "mars_field_delivery_report_final.pdf"
HTML_PATH = OUT_DIR / "mars_field_delivery_report_final.html"


def resolve_writable_path(path: Path) -> Path:
    if not path.exists():
        return path
    try:
        with path.open("ab"):
            return path
    except PermissionError:
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return path.with_name(f"{path.stem}_{stamp}{path.suffix}")


def html_escape(value: object) -> str:
    return html.escape("" if value is None else str(value))


def set_font(run, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_p(doc_or_cell, text: str, *, size: float = 10.5, bold: bool = False, italic: bool = False, align=None, color: str | None = None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    return p


def add_h(doc_or_cell, text: str, level: int = 1) -> None:
    size = {1: 15.5, 2: 12.5, 3: 11.0}[level]
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(3)


def add_bullets(doc_or_cell, items: list[str], size: float = 10.0) -> None:
    for item in items:
        p = doc_or_cell.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=size)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "D9E2F3")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_font(run, size=9, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_font(run, size=8.6)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.0) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_font(run, size=8.5, italic=True)
    cap.paragraph_format.space_after = Pt(6)


def build_print_html() -> str:
    framework_img = (ROOT / "reports" / "mars_resilience_stack" / "figures" / "figure1_three_layer_framework.png").resolve().as_uri()
    benchmark_img = (MARS_ROOT / "outputs" / "paper_bundle_v1" / "figures" / "figure2_benchmark_overview_v3.png").resolve().as_uri()
    decoder_img = (MARS_ROOT / "outputs" / "paper_bundle_v1" / "figures" / "figure3_decoder_calibration_v3.png").resolve().as_uri()
    case_img = (MARS_ROOT / "outputs" / "paper_bundle_v1" / "figures" / "figure4_case_studies_master_v3.png").resolve().as_uri()
    priority_rows = delivery.build_priority_rows()
    experiment_rows = delivery.build_next_experiments()
    runbook_rows = delivery.build_runbook_commands()[:8]

    def table(headers: list[str], rows: list[list[str]]) -> str:
        head = "".join(f"<th>{h}</th>" for h in headers)
        body = []
        for row in rows:
            body.append("<tr>" + "".join(f"<td>{html_escape(v)}</td>" for v in row) + "</tr>")
        return f"<table><thead><tr>{head}</tr></thead><tbody>{''.join(body)}</tbody></table>"

    priority_table = table(
        ["优先级", "对象组", "蛋白 / 构建体", "为什么现在做", "下一动作"],
        [[r["priority"], r["group"], r["targets"], r["why"], r["next_action"]] for r in priority_rows],
    )
    experiment_table = table(
        ["序号", "实验", "输入对象", "目标", "关键 readout", "成功判据"],
        [[r["rank"], r["experiment"], r["inputs"], r["goal"], r["readouts"], r["success"]] for r in experiment_rows],
    )
    runbook_table = table(
        ["标签", "入口", "命令", "说明"],
        [[r["tag"], r["title"], r["command"], r["desc"]] for r in runbook_rows],
    )

    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <title>MARS-FIELD Delivery Report Final</title>
  <style>
    @page {{ size: A4; margin: 18mm 16mm; }}
    body {{
      font-family: "Segoe UI", "Microsoft YaHei", sans-serif;
      color: #1f2a2f;
      line-height: 1.6;
      font-size: 11pt;
    }}
    h1,h2,h3 {{ color: #1f4e79; }}
    h1 {{ text-align: center; font-size: 24pt; margin-bottom: 6px; }}
    .sub {{ text-align: center; color: #555; margin-bottom: 24px; }}
    .section {{ margin-top: 24px; page-break-inside: avoid; }}
    .note {{
      background: #f5eee5; border-left: 4px solid #8b3d1f; padding: 10px 12px; margin: 12px 0;
    }}
    table {{ width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 9.5pt; }}
    th, td {{ border: 1px solid #d7d0c2; padding: 8px 9px; vertical-align: top; text-align: left; }}
    th {{ background: #ece3d6; }}
    figure {{ margin: 18px 0 24px; page-break-inside: avoid; }}
    img {{ max-width: 100%; display: block; margin: 0 auto; }}
    figcaption {{ text-align: center; font-size: 9pt; color: #555; margin-top: 6px; }}
    ul {{ margin-top: 8px; }}
  </style>
</head>
<body>
  <h1>MARS-FIELD 总交付报告</h1>
  <div class="sub">Final delivery package for GitHub release, HTML handoff, and manuscript support</div>

  <section class="section">
    <h2>一、交付结论</h2>
    <p>当前仓库和交付内容已经足够作为研究原型进行 GitHub 发布与论文素材交付。它已经具备统一 evidence-to-sequence 设计框架、单目标 pipeline、多目标 benchmark、paper bundle、结构案例和中文/HTML 交付文稿等完整资产。</p>
    <div class="note">
      当前最准确的定位是 <strong>MARS-FIELD engineering approximation v1</strong>。它已经是可运行的研究原型平台，但还不是 fully joint、fully learned 的最终神经场模型。
    </div>
  </section>

  <section class="section">
    <h2>二、项目三层结构</h2>
    <ul>
      <li><strong>Cld 主线：</strong> 面向火星相关多压力条件下的主功能酶路线。</li>
      <li><strong>AresG / DrwH：</strong> 作为保护增强模块和 cargo-cap 路线。</li>
      <li><strong>MARS-FIELD / MarsStack：</strong> 统一结构、进化、祖先、retrieval 和记忆、环境条件与候选解码的设计平台。</li>
    </ul>
    <figure>
      <img src="{framework_img}" alt="three layer framework" />
      <figcaption>图 1 | 项目三层统一框架。</figcaption>
    </figure>
  </section>

  <section class="section">
    <h2>三、主线蛋白优先级重排</h2>
    {priority_table}
  </section>

  <section class="section">
    <h2>四、现在最该先做的实验</h2>
    {experiment_table}
  </section>

  <section class="section">
    <h2>五、平台方法与结果总览</h2>
    <p>MARS-FIELD 当前已经实现单目标 pipeline、多目标 benchmark、family/held-out 统计、ablation、neural comparison 和 paper bundle 资产输出。其关键抽象不是单一生成器，而是共享 residue field 与 pairwise tensor。</p>
    <figure>
      <img src="{benchmark_img}" alt="benchmark overview" />
      <figcaption>图 2 | Twelvepack benchmark 总览。</figcaption>
    </figure>
    <figure>
      <img src="{decoder_img}" alt="decoder calibration" />
      <figcaption>图 3 | Decoder 与 calibration 分析。</figcaption>
    </figure>
    <figure>
      <img src="{case_img}" alt="case studies" />
      <figcaption>图 4 | 代表性案例总览。</figcaption>
    </figure>
  </section>

  <section class="section">
    <h2>六、代码入口与运行手册摘要</h2>
    {runbook_table}
    <div class="note">
      更完整的运行入口、专题脚本和交付包内容已经放在 HTML 交付目录中，供浏览器直接查看。
    </div>
  </section>

  <section class="section">
    <h2>七、GitHub release 跨机器可用性改进</h2>
    <ul>
      <li>新增 <code>requirements.txt</code>。</li>
      <li>新增 <code>docs/setup_cross_machine_release_v1.md</code>。</li>
      <li>新增 <code>scripts/check_mars_runtime.py</code>。</li>
      <li>移除多个配置里的本地 <code>python_executable</code> 硬编码，默认回退到当前 <code>sys.executable</code>。</li>
      <li>将 Cld topic/notopic 配置里的 <code>F:/...</code> 数据路径改为相对路径。</li>
      <li>将旧的打包脚本默认输出目录改成相对工作区的 <code>release_packages</code>。</li>
    </ul>
  </section>

  <section class="section">
    <h2>八、对外口径</h2>
    <ul>
      <li>可以说它是 benchmarked protein engineering research prototype。</li>
      <li>可以说它已经实现 unified evidence-to-sequence workflow。</li>
      <li>不要说它是 completed fully neural end-to-end field model。</li>
      <li>不要说它是 final production system。</li>
    </ul>
  </section>
</body>
</html>
"""


def export_pdf_via_edge(html_path: Path, pdf_path: Path) -> bool:
    edge_candidates = [
        Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
        Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
    ]
    edge = next((path for path in edge_candidates if path.exists()), None)
    if edge is None:
        return False
    cmd = [
        str(edge),
        "--headless=new",
        "--disable-gpu",
        "--no-pdf-header-footer",
        f"--print-to-pdf={pdf_path}",
        html_path.resolve().as_uri(),
    ]
    result = subprocess.run(
        cmd,
        check=False,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        timeout=120,
    )
    return result.returncode == 0 and pdf_path.exists()


def build_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    add_p(doc, "MARS-FIELD 总交付报告", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color="1F4E79")
    add_p(doc, "Final delivery package for GitHub release, HTML handoff, and manuscript support", size=12.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "当前定位：MARS-FIELD engineering approximation v1", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_h(doc, "一、交付结论", 1)
    add_p(
        doc,
        "当前仓库和交付内容已经足够作为研究原型进行 GitHub 发布与论文素材交付。它已经具备统一 evidence-to-sequence 设计框架、单目标 pipeline、多目标 benchmark、paper bundle、结构案例和中文/HTML 交付文稿等完整资产。"
    )
    add_p(
        doc,
        "但当前最准确的表述仍然是工程近似版而非 fully joint 的最终神经场模型。因此，这次总交付的目标不是夸大完成度，而是把已经成立的能力、代码组织、实验优先级、运行入口和对外口径一起收清楚。"
    )
    add_bullets(
        doc,
        [
            "可以直接用于 GitHub / 论文素材整理 / 内部汇报",
            "可以作为当前工作区下的稳定研究原型运行主要流程",
            "不应对外描述为 fully end-to-end、fully joint 的最终模型",
        ],
    )

    add_h(doc, "二、项目三层结构", 1)
    add_p(doc, "当前项目最适合被理解成三层统一结构：Cld 主功能酶路线、AresG/DrwH 保护模块路线，以及 MARS-FIELD / MarsStack 方法平台。")
    add_bullets(
        doc,
        [
            "Cld：火星相关多压力条件下的主功能酶路线，当前最值得收敛为主叙事。",
            "AresG：应激下形成可逆保护微环境的模块，更适合承担 cargo 保护和配方增强角色。",
            "DrwH：更像紧凑 cargo-cap，适合保护酶活 cargo，而不是 DNA shielding tail。",
            "MARS-FIELD / MarsStack：统一结构、进化、祖先、retrieval 和记忆、环境条件与候选解码的设计平台。",
        ],
    )
    add_figure(
        doc,
        ROOT / "reports" / "mars_resilience_stack" / "figures" / "figure1_three_layer_framework.png",
        "图 1 | 项目三层统一框架：Cld 主功能层、AresG/DrwH 保护模块层、MARS-FIELD 方法平台层。",
        width=6.2,
    )

    add_h(doc, "三、主线蛋白与构建体清单", 1)
    priority_rows = delivery.build_priority_rows()
    add_table(
        doc,
        ["优先级", "对象组", "蛋白 / 构建体", "为什么现在做", "下一动作"],
        [[row["priority"], row["group"], row["targets"], row["why"], row["next_action"]] for row in priority_rows],
        [0.6, 1.5, 1.6, 2.2, 1.8],
    )
    add_p(doc, "这张表决定了近期实验资源应该先投到哪里：先压实 Cld 主线，再让 AresG 和 DrwH 作为保护增强模块逐步接上。")

    add_h(doc, "四、现在最该先做的实验", 1)
    experiment_rows = delivery.build_next_experiments()
    add_table(
        doc,
        ["序号", "实验", "输入对象", "目标", "关键 readout", "成功判据"],
        [[row["rank"], row["experiment"], row["inputs"], row["goal"], row["readouts"], row["success"]] for row in experiment_rows],
        [0.5, 1.4, 1.5, 1.9, 2.0, 1.8],
    )
    add_p(doc, "当前最不建议的做法，是在 Cld 骨架和 stress proxy readout 还没收敛前，就把更多保护模块、更多新蛋白和更多 benchmark 一起平均推进。")

    add_h(doc, "五、方法平台与 benchmark 结果", 1)
    add_p(
        doc,
        "MARS-FIELD 当前已经实现单目标 pipeline、多目标 benchmark、family/held-out 统计、ablation、neural comparison 和 paper bundle 资产输出。它最关键的抽象不是某一个生成器，而是把结构、进化、祖先、retrieval 和记忆、环境条件统一投到共享 residue field 和 pairwise tensor 中。"
    )
    add_figure(
        doc,
        MARS_ROOT / "outputs" / "paper_bundle_v1" / "figures" / "figure1_mars_field_architecture_v2.png",
        "图 2 | MARS-FIELD 架构图：多源证据进入共享 residue field，再经 decoder、selector 与 hybrid policy 输出候选。",
        width=6.15,
    )
    add_figure(
        doc,
        MARS_ROOT / "outputs" / "paper_bundle_v1" / "figures" / "figure2_benchmark_overview_v3.png",
        "图 3 | Twelvepack benchmark 总览图：平台已经具备跨 target、跨 family 的验证能力。",
        width=6.15,
    )
    add_figure(
        doc,
        MARS_ROOT / "outputs" / "paper_bundle_v1" / "figures" / "figure3_decoder_calibration_v3.png",
        "图 4 | Decoder 与 calibration 分析图：展示 field-style 候选解码和 final policy 校准关系。",
        width=6.15,
    )

    bench_rows = delivery.build_benchmark_rows()
    add_table(
        doc,
        ["Target", "Family", "设计位点", "为什么选它", "当前 policy winner", "来源"],
        [[r["target"], r["family"], r["design_positions"], r["role"], r["overall"], r["overall_source"]] for r in bench_rows],
        [1.0, 1.0, 1.0, 2.5, 1.8, 1.0],
    )

    add_h(doc, "六、代码布局与运行入口", 1)
    add_p(doc, "当前仓库已经按照专题脚本区、设计结果区、平台代码区、报告区和交付区形成比较清晰的分层。下面列出最重要的执行入口。")
    add_table(
        doc,
        ["入口脚本", "作用", "说明"],
        [
            ["scripts/run_mars_pipeline.py", "单目标主流程", "候选生成、field 构建、排序、decoder、neural rerank 与输出。"],
            ["scripts/run_mars_benchmark.py", "多目标 benchmark", "统一跑 twelvepack / ninepack，并输出 family/held-out/ablation。"],
            ["scripts/build_structure_motif_atlas.py", "motif atlas 构建", "建立 retrieval branch 的本地结构模体记忆。"],
            ["scripts/run_mars_field_neural_reranker.py", "神经重排入口", "训练并运行 holdout neural reranker。"],
            ["scripts/build_mars_field_html_delivery.py", "HTML 交付生成", "重建本次 HTML 交付目录与 zip。"],
            ["scripts/check_mars_runtime.py", "运行环境检查", "检查核心 Python 包、外部工具与重要路径是否存在。"],
        ],
        [2.1, 1.4, 3.0],
    )

    add_h(doc, "七、GitHub release 跨机器可用性改进", 1)
    add_p(doc, "为了让仓库更适合跨机器使用，我本轮补了几项关键改进：")
    add_bullets(
        doc,
        [
            "新增 requirements.txt，声明核心 Python 依赖。",
            "新增 docs/setup_cross_machine_release_v1.md，明确基础环境、可选组件和 setup 步骤。",
            "新增 scripts/check_mars_runtime.py，帮助快速检查当前机器是否具备主要运行条件。",
            "run_mars_pipeline.py 中的 ESM-IF 解释器现在默认回退到当前 sys.executable，而不再强依赖本地 C:/Python 路径。",
            "多个 target config 中的 python_executable 硬编码已移除。",
            "CLD_3Q09_TOPIC / NOTOPIC 中原本的 F:/ 绝对数据路径已改为相对路径，避免换机器直接失效。",
            "旧的 build_mars_field_v1_package.ps1 默认打包路径已改成相对工作区 release_packages，而不是绑定本机绝对目录。",
        ],
    )
    add_p(doc, "这意味着当前仓库已经更适合做 GitHub release，但仍然不是零配置生产系统。ProteinMPNN、ESM-IF、MAFFT、IQ-TREE、AF3 等外部组件仍需要按目标用途分别配置。")

    add_h(doc, "八、HTML / Word / PDF 交付物", 1)
    add_bullets(
        doc,
        [
            f"HTML 最终目录：{ROOT / 'release_packages' / 'mars_field_html_delivery_v2_2_final'}",
            f"HTML 最终压缩包：{ROOT / 'release_packages' / 'mars_field_html_delivery_v2_2_final_20260419_114350.zip'}",
            f"正式 Word 交付稿：{DOCX_PATH}",
            f"正式 PDF 交付稿：{PDF_PATH}（若本机 Word COM 可用则自动导出）",
        ],
    )

    add_h(doc, "九、对外口径", 1)
    add_p(doc, "当前最稳妥、也最诚实的对外表述是：")
    add_bullets(
        doc,
        [
            "MARS-FIELD is a benchmarked protein engineering research prototype.",
            "It already implements a unified evidence-to-sequence workflow.",
            "It should be described as an engineering approximation v1, not as a fully learned end-to-end final model.",
        ],
    )
    add_p(doc, "如果用于 GitHub、论文素材整理和汇报，这个口径是完全成立的。如果用于宣称最终算法终点，则还需要更强的联合训练与跨机器零配置能力支持。")

    return doc


def export_pdf_via_word(docx_path: Path, pdf_path: Path) -> bool:
    powershell_script = rf"""
$ErrorActionPreference = "Stop"
$word = $null
$doc = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $doc = $word.Documents.Open("{docx_path}")
    $doc.SaveAs([ref] "{pdf_path}", [ref] 17)
    $doc.Close()
    $word.Quit()
}} catch {{
    if ($doc -ne $null) {{ $doc.Close() }}
    if ($word -ne $null) {{ $word.Quit() }}
    throw
}}
"""
    result = subprocess.run(
        ["powershell", "-NoProfile", "-Command", powershell_script],
        check=False,
        capture_output=True,
        text=True,
    )
    return result.returncode == 0 and pdf_path.exists()


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = resolve_writable_path(DOCX_PATH)
    html_path = resolve_writable_path(HTML_PATH)
    pdf_path = resolve_writable_path(PDF_PATH)
    doc = build_doc()
    doc.save(docx_path)
    print(f"Saved delivery Word report to {docx_path}")
    html_path.write_text(build_print_html(), encoding="utf-8")
    print(f"Saved delivery HTML report to {html_path}")
    if export_pdf_via_edge(html_path, pdf_path):
        print(f"Saved delivery PDF report to {pdf_path}")
    elif export_pdf_via_word(docx_path, pdf_path):
        print(f"Saved delivery PDF report to {pdf_path}")
    else:
        print("PDF export skipped or failed; Edge/Word automation not available.")


if __name__ == "__main__":
    main()
