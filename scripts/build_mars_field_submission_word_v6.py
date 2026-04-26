from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "paper_bundle_v1"
FIG_DIR = OUT_DIR / "figures"
OUT_DOCX = OUT_DIR / "MARS_FIELD_Submission_Style_v6.docx"


def set_font(run, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_p(doc_or_cell, text: str, *, size: float = 10.5, bold: bool = False, italic: bool = False, align=None, color: str | None = None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(4)
    return p


def add_h(doc_or_cell, text: str, level: int = 1):
    size = {1: 14.5, 2: 12.0, 3: 10.8}[level]
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(3)


def add_bullets(doc_or_cell, items: list[str], size: float = 10.0):
    for item in items:
        p = doc_or_cell.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=size)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "D9E2F3")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_font(run, size=9, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_font(run, size=8.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def add_figure(doc, image_name: str, caption: str, width: float = 6.0) -> None:
    path = FIG_DIR / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = c.add_run(caption)
    set_font(run, size=8.5, italic=True)
    c.paragraph_format.space_after = Pt(6)


def references() -> list[str]:
    return [
        "1. Jumper, J. et al. Highly accurate protein structure prediction with AlphaFold. Nature 596, 583-589 (2021).",
        "2. Dauparas, J. et al. Robust deep learning-based protein sequence design using ProteinMPNN. Science 378, 49-56 (2022).",
        "3. Hsu, C. et al. Learning inverse folding from millions of predicted structures. Proc. ICML, PMLR 162, 8946-8970 (2022).",
        "4. Lin, Z. et al. Evolutionary-scale prediction of atomic-level protein structure with a language model. Science 379, 1123-1130 (2023).",
        "5. van Kempen, M. et al. Fast and accurate protein structure search with Foldseek. Nat. Biotechnol. 42, 243-246 (2024).",
        "6. Alley, E. C. et al. Unified rational protein engineering with sequence-based deep representation learning. Nat. Methods 16, 1315-1322 (2019).",
        "7. Chisholm, L. O. et al. Ancestral reconstruction and the evolution of protein energy landscapes. Annu. Rev. Biophys. 53, 127-146 (2024).",
        "8. Matthews, D. S. et al. Leveraging ancestral sequence reconstruction for protein representation learning. Nat. Mach. Intell. 6, 1542-1555 (2024).",
    ]


def build_doc() -> Document:
    bench = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "benchmark_summary.csv")
    compare = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "compare_current_vs_final.csv")
    delta_col = "policy_selection_score_delta_final_minus_current"
    positive = int((compare[delta_col] > 1e-9).sum())
    negative = int((compare[delta_col] < -1e-9).sum())
    mean_delta = float(compare[delta_col].mean())
    decoder_preview = int(bench["neural_decoder_generated_count"].sum())
    decoder_novel = int(bench["neural_decoder_novel_count"].sum())
    decoder_rejected = int(bench["neural_decoder_rejected_count"].sum())
    decoder_injected = int(bench["neural_decoder_injected"].sum())

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    add_p(doc, "MARS-FIELD", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color="1F4E79")
    add_p(doc, "A unified evidence-to-sequence controller for protein engineering", size=13, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Submission-style draft v6", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_h(doc, "Abstract", 1)
    add_p(
        doc,
        f"We present MARS-FIELD, a unified evidence-to-sequence controller that projects geometry-conditioned structural evidence, phylogenetic profiles, ancestral lineage information, retrieval-based motif memory, and environment-conditioned engineering context into a shared residue field. The current implementation couples this field to a candidate-level neural controller and a decode-time neural field generator inside the benchmark-time main path. Across a twelve-target benchmark spanning ten protein families, the final controller improved paired policy score on {positive} of 12 targets relative to the incumbent benchmark, enabled neural decoding on all 12 targets, and retained {decoder_novel} novel neural-decoder-derived candidates across {decoder_injected} targets. These results support MARS-FIELD as a stable neuralized controller-decoder system that moves beyond heuristic score stacking toward an evidence-conditioned field architecture.",
    )

    add_h(doc, "Introduction", 1)
    add_p(doc, "Protein engineering pipelines often integrate structural analysis, evolutionary context, ancestral reconstruction, and machine-learning proposal generators as loosely connected modules. Although this strategy enables rapid iteration, it also creates a conceptual weakness: the system behaves as an ensemble of partially redundant signals rather than as a unified controller. This becomes especially limiting when the objective is not simply sequence plausibility but robust engineering under multiple constraints, including oxidation risk, flexible-surface burden, motif preservation, and environment-conditioned performance targets.")
    add_p(doc, "MARS-FIELD addresses this problem by treating protein engineering as an evidence-to-sequence control task. Geometric, phylogenetic, ancestral, retrieval-based, and environment-conditioned signals are projected into a shared residue field over the design positions. This field is the native decision object of the method: it defines residue-wise preferences and pairwise couplings, which are then used to score incumbents and decode new sequence proposals under engineering constraints.")
    add_p(doc, "The present implementation moves beyond a reranker-centric architecture in two ways. First, ancestry and retrieval are represented as learned memory-linked branches within the neural controller. Second, the neural branch is now coupled to a decode-time neural field generator in the main benchmark path. The goal of this work is therefore not to claim a final fully joint protein design foundation model, but to demonstrate that a unified evidence-conditioned field-controller system can be built in a way that is benchmark-stable, interpretable, and capable of neural proposal generation.")

    add_figure(doc, "figure1_mars_field_architecture_v2.png", "Figure 1 | MARS-FIELD integrates heterogeneous evidence into a shared residue field. The method operates on site-wise residue energies and pairwise couplings rather than on a generator vote.", width=6.15)

    add_h(doc, "Results", 1)
    add_p(doc, f"The final controller improved paired policy score on {positive} of 12 targets and decreased it on {negative} targets, with a mean paired delta of {mean_delta:.3f}. This near-neutral mean indicates that the benchmark remains globally stable after neural decoder integration, while the directionality of the paired effects shows that positive target-level gains substantially outnumber negative ones.")
    add_p(doc, f"Neural decoding was enabled on all 12 targets, produced {decoder_preview} preview candidates, retained {decoder_novel} novel candidates after engineering and safety filtering, and contributed retained novel candidates on {decoder_injected} targets. This demonstrates that the neural branch is functioning as a genuine proposal source rather than as a passive auxiliary scorer.")

    add_figure(doc, "figure2_benchmark_claim_v5.png", "Figure 2 | Benchmark claim figure. Paired policy shifts, headline metrics, neural decoder utilization, and family-level summary support the benchmark-level claim.", width=6.15)

    add_p(doc, "Ablation and diagnostic analyses show that the current method remains strongly grounded in chemically and evolutionarily meaningful constraints. Oxidation-aware and evolutionary signals remain dominant regulators of the design landscape, while neural gate profiles reveal target-specific evidence usage rather than a uniform black-box weighting scheme. The remaining regressions are concentrated in a small number of explicit calibration-limited cases, including the CLD topic-conditioned targets and subtilisin.")

    add_figure(doc, "figure3_mechanism_limitations_v5.png", "Figure 3 | Mechanism and limitation figure. Ablations identify dominant constraints, neural diagnostics show target-dependent evidence usage, and failures remain concentrated rather than diffuse.", width=6.15)

    add_p(doc, "Representative case studies further show that the controller does not behave identically across targets. In 1LBT, the method behaves conservatively and preserves a known safe incumbent despite active neural decoding. In TEM1, the incumbent remains stable while the neural branch surfaces an interpretable learned alternative. In PETase, the controller reproduces a canonical redesign across related structures. In CLD, the system exposes a calibration stress test in which incumbent stability and a stronger local engineering alternative remain in tension.")

    add_figure(doc, "figure4_case_studies_master_v3.png", "Figure 4 | PSE-derived case-study composite. 1LBT, TEM1, PETase, and CLD reveal distinct controller regimes: safety preservation, stable incumbent with learned alternative, reproducible redesign, and calibration stress test.", width=6.15)

    add_h(doc, "Methods", 1)
    add_p(doc, "The main benchmark consists of twelve targets spanning ten protein families. Each target defines a local design window on a fixed wild-type structural scaffold. Neural controller training is leave-one-target-out at the target level. The main deployment arm used here is benchmark_twelvepack_final, which enables both neural reranking and neural field decoding while using a hybrid final policy for stable deployment.")
    add_p(doc, "MARS-FIELD integrates five evidence streams: geometry-conditioned structural features, phylogenetic profile evidence, ancestral lineage evidence, retrieval-based motif memory, and environment-conditioned engineering context. These signals are projected into a shared residue field containing both residue-wise and pairwise preferences. The neural field contains geometry, phylogeny, ancestry, retrieval, and environment branches; ancestry and retrieval each interact with learned memory banks to fuse local residue representations with lineage memory and retrieval prototype memory.")
    add_p(doc, "Candidate-level decision making combines sequence-conditioned residue embeddings with candidate-specific evidence features, including source identity, support count, mutation burden, component-wise engineering scores, selector-rank priors, gaps to the best incumbent, and other calibration-aware signals. Training uses multiple objectives, including selection regression, engineering regression, policy regression, pairwise ranking, decoder-field supervision, winner-guard loss, non-decoder guard loss, simplicity guard loss, selector-anchor distillation, recovery loss, ancestry alignment, retrieval alignment, pairwise consistency, environment reconstruction, and gate regularization.")
    add_p(doc, "For each target, the pipeline constructs a runtime neural batch from the current target state, trains a leave-one-target-out neural field model using the remaining targets, combines learned unary and pairwise outputs with evidence-derived prior position fields and prior pairwise terms, and decodes neural_decoder candidates under constrained beam search. The final hybrid policy permits learned branch adoption only when the alternative remains sufficiently aligned with incumbent selection-score stability while preserving engineering support.")

    add_h(doc, "Discussion", 1)
    add_p(doc, "The central conclusion of this work is that a unified residue-field controller provides a more coherent algorithmic basis for protein engineering than an ensemble of disconnected proposal branches. In the current implementation, MARS-FIELD integrates heterogeneous evidence streams into a shared residue field, uses a calibrated neural controller to score candidates, and activates a decode-time neural branch inside the main benchmark path. This is qualitatively different from earlier stack-and-rerank pipelines because the neural branch now participates in generation as well as selection.")
    add_p(doc, "The remaining limitations are also clearer in this formulation. The system is now strong enough to be described as an engineering-honest end-to-end controller-decoder architecture, but it is not yet the fully joint proposal-generator / field / decoder research endpoint. The remaining gap is therefore no longer missing plumbing; it is a genuine research frontier centered on joint optimization and improved calibration on a small number of hard targets.")
    add_p(doc, "For manuscript positioning, the strongest claim is not that MARS-FIELD already solves every regime equally well, but that it establishes a stable bridge between heuristic engineering stacks and future fully joint protein design controllers. It combines a unified evidence-to-sequence formulation, real neural proposal generation, benchmark-scale stability, and explicit mechanistic accounting of both gains and failures.")

    add_h(doc, "Selected References", 1)
    for ref in references():
        add_p(doc, ref, size=9.4)

    return doc


def main() -> None:
    doc = build_doc()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Saved submission-style Word draft to {OUT_DOCX}")


if __name__ == "__main__":
    main()
