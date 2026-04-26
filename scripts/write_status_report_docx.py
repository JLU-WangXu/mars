from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)


def add_title(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(10.5)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    document.add_paragraph("")


def fmt(value: object) -> str:
    if pd.isna(value):
        return "NA"
    if isinstance(value, float):
        return f"{value:.3f}".rstrip("0").rstrip(".")
    return str(value)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--ninepack-summary", type=Path, required=True)
    parser.add_argument("--ninepack-family", type=Path, required=True)
    parser.add_argument("--family-prior-compare", type=Path, required=True)
    parser.add_argument("--outpath", type=Path, required=True)
    args = parser.parse_args()

    ninepack = pd.read_csv(args.ninepack_summary)
    family_summary = pd.read_csv(args.ninepack_family)
    family_prior = pd.read_csv(args.family_prior_compare)

    doc = Document()
    set_default_font(doc)
    add_title(
        doc,
        "MarsStack Current Execution Report",
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}",
    )

    doc.add_heading("1. Current Stage", level=1)
    add_bullets(
        doc,
        [
            "Current baseline sixpack state has been snapshotted and preserved for rollback.",
            "MarsStack now includes two learned generators: ProteinMPNN and ESM-IF.",
            "A family-conditioned positive/negative evolution prior has been implemented and is live in the main pipeline.",
            "The benchmark panel has been expanded from six proteins to nine proteins.",
            "A family-prior ablation benchmark and a held-out-family unit summary have both been generated.",
        ],
    )

    doc.add_heading("2. Code Upgrades Executed", level=1)
    add_bullets(
        doc,
        [
            "Added differential family evolution utilities to marsstack/evolution.py.",
            "Extended marsstack/mars_score.py so the evolution term can include a positive-vs-negative family prior.",
            "Updated run_mars_pipeline.py so family manifests can influence both scoring and local proposal generation.",
            "Updated run_mars_benchmark.py so summaries now record family-prior metadata and support --reuse-existing.",
            "Added new benchmark targets and configs: ADK / 1S3G, esterase / 7B4Q, Mn-SOD / 1Y67.",
            "Added with-family-prior and without-family-prior comparison configs for the new family-backed targets.",
        ],
    )

    doc.add_heading("3. Ninepack Benchmark Summary", level=1)
    doc.add_paragraph(
        f"The current benchmark contains {len(ninepack)} targets, of which "
        f"{int(ninepack['family_prior_enabled'].sum())} use the new family-conditioned prior."
    )
    add_table(
        doc,
        [
            "Target",
            "Overall Winner",
            "Overall Score",
            "Best Learned",
            "Best Learned Score",
            "Family Prior",
        ],
        [
            [
                fmt(row["target"]),
                f"{fmt(row['overall_source'])}: {fmt(row['overall_mutations'])}",
                fmt(row["overall_score"]),
                f"{fmt(row['best_learned_source'])}: {fmt(row['best_learned_mutations'])}",
                fmt(row["best_learned_score"]),
                fmt(row["family_dataset_id"]),
            ]
            for _, row in ninepack.iterrows()
        ],
    )

    doc.add_heading("4. Family Prior Comparison", level=1)
    doc.add_paragraph(
        "This table compares the three family-backed targets with and without the positive/negative family prior. "
        "The goal is to isolate the added value of the differential family signal beyond the ordinary homolog profile."
    )
    add_table(
        doc,
        [
            "Target",
            "Overall Delta",
            "Best Learned Delta",
            "With Prior Overall",
            "Without Prior Overall",
            "With Prior Best Learned",
            "Without Prior Best Learned",
        ],
        [
            [
                fmt(row["target"]),
                fmt(row["overall_score_delta"]),
                fmt(row["best_learned_score_delta"]),
                fmt(row["overall_mutations_with"]),
                fmt(row["overall_mutations_without"]),
                fmt(row["best_learned_mutations_with"]),
                fmt(row["best_learned_mutations_without"]),
            ]
            for _, row in family_prior.iterrows()
        ],
    )

    doc.add_heading("5. Held-Out Family View", level=1)
    doc.add_paragraph(
        "The current ninepack is interpreted as a held-out family panel because each benchmark target corresponds to a distinct family unit."
    )
    add_table(
        doc,
        ["Family", "N Targets", "Mean Overall Score", "Mean Best Learned Score"],
        [
            [
                fmt(row["family"]),
                fmt(row["n_targets"]),
                fmt(row["mean_overall_score"]),
                fmt(row["mean_best_learned_score"]),
            ]
            for _, row in family_summary.iterrows()
        ],
    )

    doc.add_heading("6. My Evaluation", level=1)
    add_bullets(
        doc,
        [
            "The project has moved beyond a single-protein prototype and now qualifies as a method-style benchmark stack.",
            "The most important algorithmic upgrade in this round is not ESM-IF itself, but that evolution prior is now part of the main executable pipeline rather than a placeholder idea.",
            "The family prior produces a real measurable gain for ADK and a meaningful best-learned improvement for SOD, while esterase is close to neutral. This is a healthy result: the prior is contributing signal, but not inflating every target artificially.",
            "The chemistry-aware local branch is still too strong relative to the learned branches on many targets. That means the current bottleneck remains proposal quality, not reranking.",
            "The current held-out-family view is useful for reporting, but still weak as a publication-grade generalization claim because each family currently has only one benchmark target.",
            "The next best step is to formalize a family-split benchmark protocol and add more targets per family category or more family categories with matched assay logic.",
        ],
    )

    doc.add_heading("7. Recommendation", level=1)
    add_bullets(
        doc,
        [
            "Keep the current ninepack as the new working benchmark baseline.",
            "Preserve the family prior as a permanent optional branch in the method.",
            "Next implement a stricter held-out-family benchmark protocol plus explicit with-prior / without-prior reporting in the main result tables.",
            "After that, prioritize stronger learned proposal generation rather than more hand-tuned score terms.",
        ],
    )

    doc.add_heading("8. Key Output Files", level=1)
    add_bullets(
        doc,
        [
            r"outputs\benchmark_ninepack\benchmark_summary.csv",
            r"outputs\benchmark_ninepack\family_summary.csv",
            r"outputs\benchmark_ninepack\heldout_family_units.md",
            r"outputs\benchmark_family_prior_compare\family_prior_comparison.csv",
            r"outputs\benchmark_ninepack\algorithm_upgrade_summary_2026-04-15.md",
            r"outputs\snapshots\benchmark_sixpack_2026-04-15_pre_v1",
        ],
    )

    args.outpath.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.outpath)
    print(f"Wrote Word report to {args.outpath}")


if __name__ == "__main__":
    main()
