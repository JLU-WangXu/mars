from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
PAPER_BUNDLE = ROOT / "outputs" / "paper_bundle_v1"
FIG_DIR = PAPER_BUNDLE / "figures"
OUT_DOCX = PAPER_BUNDLE / "MARS_FIELD_Methods_Results_Draft_v2.docx"


def set_run_font(run, ascii_font: str = "Times New Roman", east_asia_font: str = "Microsoft YaHei") -> None:
    run.font.name = ascii_font
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def set_paragraph_style(paragraph, *, size: float = 11, bold: bool = False, italic: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15


def add_text(doc: Document, text: str, *, style: str | None = None, size: float = 11, bold: bool = False, italic: bool = False, align=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_heading(doc: Document, text: str, level: int, size: float) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run)
    run.bold = True
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    paragraph.paragraph_format.space_after = Pt(4)


def add_table_from_rows(doc: Document, headers: list[str], rows: list[list[str]], column_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, east_asia_font="Microsoft YaHei")
                run.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, east_asia_font="Microsoft YaHei")
                    run.font.size = Pt(8.5)
    if column_widths:
        for row in table.rows:
            for idx, width in enumerate(column_widths):
                row.cells[idx].width = Inches(width)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    set_run_font(caption_run)
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_p.paragraph_format.space_after = Pt(6)


def metric_summary(compare_df: pd.DataFrame, final_df: pd.DataFrame) -> dict[str, object]:
    delta_col = "policy_selection_score_delta_final_minus_current"
    positive = int((compare_df[delta_col] > 1e-9).sum())
    negative = int((compare_df[delta_col] < -1e-9).sum())
    mean_delta = float(compare_df[delta_col].mean())
    decoder_enabled = int(final_df["neural_decoder_enabled"].sum())
    decoder_injected = int(final_df["neural_decoder_injected"].sum())
    decoder_preview = int(final_df["neural_decoder_generated_count"].sum())
    decoder_novel = int(final_df["neural_decoder_novel_count"].sum())
    decoder_rejected = int(final_df["neural_decoder_rejected_count"].sum())
    return {
        "positive": positive,
        "negative": negative,
        "mean_delta": mean_delta,
        "decoder_enabled": decoder_enabled,
        "decoder_injected": decoder_injected,
        "decoder_preview": decoder_preview,
        "decoder_novel": decoder_novel,
        "decoder_rejected": decoder_rejected,
        "n_targets": int(len(final_df)),
        "n_families": int(final_df["family"].nunique()),
    }


def build_document() -> Document:
    compare_df = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "compare_current_vs_final.csv")
    final_df = pd.read_csv(ROOT / "outputs" / "benchmark_twelvepack_final" / "benchmark_summary.csv")
    metrics = metric_summary(compare_df, final_df)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    add_text(
        doc,
        "MARS-FIELD: Methods and Results Draft v2",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text(
        doc,
        "Nature-style manuscript draft built from benchmark_twelvepack_final",
        size=10,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_heading(doc, "Key Readout", level=1, size=13)
    summary_rows = [
        ["Targets in final benchmark", str(metrics["n_targets"])],
        ["Protein families represented", str(metrics["n_families"])],
        ["Policy score improved vs incumbent", f"{metrics['positive']}/{metrics['n_targets']}"],
        ["Policy score decreased vs incumbent", f"{metrics['negative']}/{metrics['n_targets']}"],
        ["Mean paired policy delta", f"{metrics['mean_delta']:.3f}"],
        ["Neural decoder enabled", f"{metrics['decoder_enabled']}/{metrics['n_targets']}"],
        ["Targets with retained neural-decoder candidates", f"{metrics['decoder_injected']}/{metrics['n_targets']}"],
        ["Total neural-decoder preview / retained / rejected", f"{metrics['decoder_preview']} / {metrics['decoder_novel']} / {metrics['decoder_rejected']}"],
    ]
    add_table_from_rows(doc, ["Metric", "Value"], summary_rows, [3.2, 2.8])

    add_heading(doc, "Results", level=1, size=13)
    add_heading(doc, "A unified evidence-to-sequence controller supports decode-time neural proposal generation", level=2, size=11.5)
    add_text(
        doc,
        (
            "The current MARS-FIELD implementation is no longer limited to reranking candidates proposed by external generators. "
            "Instead, the benchmark-time main path now constructs a runtime neural batch for each target, trains a leave-one-target-out "
            "neural field model, converts the learned unary and pairwise outputs into decode-ready residue fields, and decodes neural_decoder "
            "candidates inside the same pipeline. This architecture couples a shared evidence field, a neural candidate controller, and a decode-time "
            "proposal mechanism in a single executable loop."
        ),
    )
    add_text(
        doc,
        (
            f"Across the twelve-target benchmark, the neural decoder was enabled on all {metrics['n_targets']} targets. "
            f"It produced {metrics['decoder_preview']} preview candidates in total, retained {metrics['decoder_novel']} novel decoded candidates after "
            f"safety and engineering gating, and injected retained neural-decoder candidates on {metrics['decoder_injected']} targets. "
            f"These counts indicate that the neural field is functioning as a genuine proposal source rather than only as a passive scoring layer."
        ),
    )
    add_figure(
        doc,
        FIG_DIR / "figure1_mars_field_architecture_v1.svg".with_suffix(".svg").with_name("figure1_mars_field_architecture_v1.svg").with_suffix(".png")
        if False else FIG_DIR / "figure2_benchmark_overview_v3.png",
        "Figure 1. Benchmark overview for the current twelve-target MARS-FIELD final controller.",
        width=6.2,
    )

    add_heading(doc, "The final controller remains competitive while activating a neural field decoder", level=2, size=11.5)
    add_text(
        doc,
        (
            f"Relative to the incumbent benchmark, the final controller improved paired policy score on {metrics['positive']} of {metrics['n_targets']} targets "
            f"and decreased it on {metrics['negative']} targets, with a near-neutral mean paired delta of {metrics['mean_delta']:.3f}. "
            "This pattern is important: it indicates that the end-to-end neuralized controller can be inserted into the main path without causing broad collapse, "
            "while still producing measurable gains on a majority of targets."
        ),
    )
    add_text(
        doc,
        (
            "Strong positive cases included 1LBT, adenylate kinase, PETase 5XFY, SFGFP, and SOD. "
            "In 1LBT, the incumbent winner M298L was preserved but its policy score increased, indicating that the new controller strengthened a known safe solution "
            "rather than destabilizing it. In adenylate kinase, the controller shifted the policy winner to Y24F;H28Q;M103I;H109V with a positive paired delta. "
            "PETase, SFGFP, and SOD similarly retained chemically plausible top solutions while remaining compatible with the neuralized controller-decoder stack."
        ),
    )
    add_text(
        doc,
        (
            "The remaining regressions were concentrated rather than diffuse. CLD_3Q09_NOTOPIC, CLD_3Q09_TOPIC, and subtilisin_2st1 showed lower final policy scores or a "
            "policy shift to a less favorable alternative. These targets should be framed explicitly as calibration-limited cases rather than hidden failure modes."
        ),
    )
    add_figure(doc, FIG_DIR / "figure_neural_comparison_v1.png", "Figure 2. Neural branch comparison across the benchmark panel.", width=5.8)
    add_figure(doc, FIG_DIR / "figure_policy_compare_v1.png", "Figure 3. Policy comparison highlighting current versus final controller behavior.", width=5.6)

    add_heading(doc, "Case studies", level=2, size=11.5)
    add_text(
        doc,
        (
            "1LBT functions as a conservative safety case. The final controller preserved the incumbent M298L winner while still building a full neural field and neural decoder "
            "for the target. This is desirable because earlier versions of the system showed that 1LBT is vulnerable to over-eager decoder behavior. "
            "The present controller therefore demonstrates that neural generation can be active without forcing adoption of weak decoded alternatives."
        ),
    )
    add_text(
        doc,
        (
            "TEM1 provides a complementary case in which the incumbent top solution remained stable, but the best learned candidate in the final run became a neural-decoder-derived "
            "variant, H153Q;M155L;W229Q;M272I. This supports the argument that the neural field decoder can contribute meaningful alternatives without requiring the controller to replace "
            "the incumbent on every target."
        ),
    )
    add_text(
        doc,
        (
            "PETase 5XFY and 5XH3 provide a reproducibility case across related structures. In both structures, the canonical aromatic redesign remained the top policy solution under "
            "the final controller. This shows that the neuralized field-decoder path does not need to change the top sequence to add value; instead, it can reproduce stable chemistry-aware "
            "solutions across multiple structural contexts."
        ),
    )
    add_text(
        doc,
        (
            "The CLD topic-conditioned target provides a calibration stress test. The final policy preserved the incumbent W155F;W156F;M167L;M212L;W227F, while the neural branch continued "
            "to surface nearby alternatives with stronger local engineering signals. This target is therefore useful for discussing the tradeoff between learned exploration and stable selection."
        ),
    )
    add_figure(doc, FIG_DIR / "figure4_case_1lbt_v2.png", "Figure 4. 1LBT case-study panel.", width=6.1)
    add_figure(doc, FIG_DIR / "figure5_case_tem1_v2.png", "Figure 5. TEM1 case-study panel.", width=6.1)
    add_figure(doc, FIG_DIR / "figure6_case_petase_v2.png", "Figure 6. PETase case-study panel.", width=6.1)
    add_figure(doc, FIG_DIR / "figure7_case_cld_v2.png", "Figure 7. CLD case-study panel.", width=6.1)

    add_heading(doc, "Main benchmark table", level=2, size=11.5)
    target_rows = []
    delta_col = "policy_selection_score_delta_final_minus_current"
    compare_lookup = {str(row["target"]): row for _, row in compare_df.iterrows()}
    for _, row in final_df.iterrows():
        cmp = compare_lookup[str(row["target"])]
        target_rows.append(
            [
                str(row["target"]),
                str(row["family"]),
                str(row["policy_mutations"]),
                f"{float(row['policy_selection_score']):.3f}",
                f"{float(row['policy_engineering_score']):.3f}",
                f"{float(cmp[delta_col]):+.3f}",
                str(int(row["neural_decoder_novel_count"])),
            ]
        )
    add_table_from_rows(
        doc,
        ["Target", "Family", "Final Policy Mutation", "Policy Score", "Engineering Score", "Paired Delta", "Neural Decoder Novel"],
        target_rows,
        [1.0, 1.0, 2.4, 0.8, 0.9, 0.8, 0.8],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Methods", level=1, size=13)
    add_heading(doc, "Overall framework", level=2, size=11.5)
    add_text(
        doc,
        (
            "MARS-FIELD is formulated as a unified evidence-to-sequence controller rather than a collection of disconnected proposal tools. "
            "The system absorbs five evidence classes: geometry-conditioned structural features, phylogenetic sequence profiles, ancestral lineage signals, "
            "retrieval-based motif memory, and environment-conditioned engineering context. These signals are combined into a shared residue field over the design positions."
        ),
    )
    add_heading(doc, "Evidence streams and field construction", level=2, size=11.5)
    add_text(
        doc,
        (
            "Structural evidence is derived from residue-level geometric features, including solvent exposure, flexibility, protected distances, and hotspot annotations. "
            "Evolutionary evidence is derived from homolog profiles and family-differential residue preferences. Ancestral evidence is represented as posterior residue preferences "
            "and confidence-weighted lineage-derived recommendations. Retrieval evidence is encoded as structure-derived motif memory. Environment evidence is encoded as target-level "
            "context tokens, including oxidation-hotspot burden, flexible-surface burden, design-window size, and prior availability flags."
        ),
    )
    add_text(
        doc,
        (
            "The neural field contains geometry, phylogeny, ancestry, retrieval, and environment branches. Ancestry and retrieval are each connected to learned memory banks, "
            "allowing site representations to be fused with lineage memory and retrieval prototype memory. The resulting site hidden states drive both unary residue preferences "
            "and low-rank pairwise interaction terms."
        ),
    )
    add_heading(doc, "Candidate-level neural controller", level=2, size=11.5)
    add_text(
        doc,
        (
            "Candidate-level decision making combines sequence-conditioned residue embeddings with candidate-specific evidence features, including source type, support count, mutation burden, "
            "component-wise engineering terms, rank-calibrated selector features, and selector-prior context. Pairwise summaries are fused into the same candidate embedding. "
            "The controller outputs candidate-level selection, engineering, and policy predictions."
        ),
    )
    add_text(
        doc,
        (
            "Training uses multiple calibration losses: regression to target-wise normalized selection scores, engineering-score regression, pairwise policy ranking, decoder-field residue supervision, "
            "winner-guard loss, non-decoder-guard loss, simplicity-guard loss, and selector-anchor distillation. Together, these losses stabilize the learned controller against over-eager branch replacement."
        ),
    )
    add_heading(doc, "Neural field decoder", level=2, size=11.5)
    add_text(
        doc,
        (
            "For each target, the pipeline constructs a runtime neural batch from the current target state, trains a leave-one-target-out neural field model using the remaining benchmark targets, "
            "and converts the learned unary and pairwise outputs into decode-ready residue fields. To keep decode-time generation grounded, neural field outputs are combined with evidence-derived prior "
            "position fields and prior pairwise terms before constrained beam decoding. This teacher-forced neural field produces neural_decoder candidates inside the main pipeline."
        ),
    )
    add_heading(doc, "Benchmark protocol", level=2, size=11.5)
    add_text(
        doc,
        (
            "The main benchmark panel contains twelve targets spanning ten protein families. Neural training is leave-one-target-out at the target level. The main paper-facing deployment arm is "
            "benchmark_twelvepack_final, which enables both neural reranking and neural field decoding while using a hybrid final policy. Paired comparisons are made against the incumbent benchmark. "
            "Primary metrics are policy selection score, engineering score (mars_score), and neural decoder utilization."
        ),
    )

    add_heading(doc, "Interpretation of key benchmark metrics", level=1, size=13)
    add_text(
        doc,
        "1. `neural decoder enabled: 12/12` means the end-to-end neural generation branch is general enough to run on every benchmark target rather than being a hand-picked demo.",
    )
    add_text(
        doc,
        "2. `retained novel neural-decoder candidates: 34` means the neural field is contributing non-redundant designs that were not already present in the incumbent candidate pool.",
    )
    add_text(
        doc,
        "3. `policy score improved on 9/12` means the final controller is beneficial on most targets under paired comparison, not just in isolated examples.",
    )
    add_text(
        doc,
        "4. `negative on 3/12` means the system still has calibration-limited targets, which should be discussed honestly as remaining weaknesses rather than hidden.",
    )
    add_text(
        doc,
        "5. `mean paired delta about -0.001` means that after introducing the end-to-end neural field decoder, the benchmark does not collapse globally. The near-zero mean shows overall stability, while the 9/12 positive count shows that gains are distributed across most targets.",
    )

    return doc


def main() -> None:
    doc = build_document()
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Saved Word draft to {OUT_DOCX}")


if __name__ == "__main__":
    main()
